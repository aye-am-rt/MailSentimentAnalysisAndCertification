import os
import re
from docx import Document
import datetime


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def generate_certificate(d):
    name = d['receiver']
    regex1 = re.compile(r"Name")
    replace1 = name
    # date ="10/12/2019"
    date = datetime.date.today().strftime("%B %d, %Y")
    appreciator = d['sender']
    regex2 = re.compile(r"ddd")
    replace2 = date
    regex3 = re.compile(r"sss")
    replace3 = appreciator
    # os.chdir('../../../../../resources')
    print("7777777777777777777", os.getcwd())
    certificate_template = "a.docx"
    doc = Document(certificate_template)
    docx_replace_regex(doc, regex1, replace1)
    docx_replace_regex(doc, regex2, replace2)
    docx_replace_regex(doc, regex3, replace3)
    doc.save('final_certificate.docx')
    print("certificate saved in final_certificate.docx file")
